from bs4 import BeautifulSoup
import requests
import re
import docx
from docx import Document



def indeed_job_search(num_pages, position, document, cur_index):
    
    if position == 'da':    
        url = 'https://il.indeed.com/jobs?q=data+analyst&l=%D7%AA%D7%9C+%D7%90%D7%91%D7%99%D7%91+-%D7%99%D7%A4%D7%95%2C+%D7%9E%D7%97%D7%95%D7%96+%D7%AA%D7%9C+%D7%90%D7%91%D7%99%D7%91&ts=1592342093422&pts=1591532833716&rq=1&rsIdx=0'
    elif position == 'ds':
        url = 'https://il.indeed.com/jobs?q=Data+Scientist&l=%D7%AA%D7%9C+%D7%90%D7%91%D7%99%D7%91+-%D7%99%D7%A4%D7%95%2C+%D7%9E%D7%97%D7%95%D7%96+%D7%AA%D7%9C+%D7%90%D7%91%D7%99%D7%91&ts=1591532964826&rq=1&rsIdx=2&fromage=last&newcount=6'
    elif position == 'de':
        url = 'https://il.indeed.com/jobs?q=data+engineer&l=%D7%AA%D7%9C+%D7%90%D7%91%D7%99%D7%91+-%D7%99%D7%A4%D7%95%2C+%D7%9E%D7%97%D7%95%D7%96+%D7%AA%D7%9C+%D7%90%D7%91%D7%99%D7%91&ts=1592341157736&pts=1592321207153&rq=1&rsIdx=1'
        
    document.add_paragraph('www.indeed.co.il:')
    job_names, job_links = [], []
    
    for page in range(1,num_pages):
        if page == 1:
            url2 = url
        elif page > 1:
            url2 = url + '&start=' + str((page-1)*10)
        
        page = requests.get(url2)
        soup = BeautifulSoup(page.content, 'html.parser')
        for each in soup.find_all('h2'):
            name = beautify_text(each.get_text())
            job_names.append(name)
            for a in each.find_all('a', href=True):
                link = 'https://il.indeed.com' + a['href']
                job_links.append(link)

                #writing to file:
                write_to_file(document, cur_index, name, link)
            
            #increasing cur_index:
            cur_index += 1
    
    num_results = len(job_names)
        
    return num_results, job_names, job_links



def drushim_job_search(num_pages, position, document, cur_index):
    
    if position == 'da':    
        url = 'https://www.drushim.co.il/jobs/subcat/581/'
    elif position == 'ds':
        url = 'https://www.drushim.co.il/jobs/subcat/511/'
    elif position == 'de':
        url = 'https://www.drushim.co.il/jobs/subcat/582/'
        
    document.add_paragraph('www.drushim.co.il:')
    job_names, job_links = [], []
        
    for page in range(1,num_pages+1):
        if page == 1:
            url2 = url
        elif page > 1:
            url2 = url + '?page=' + str(page-1)
        
        page = requests.get(url2)
        soup = BeautifulSoup(page.content, 'html.parser')
        for each in soup.find_all('h2'):
            name = beautify_text(each.get_text())
            link = url2
            job_names.append(name) 
            job_links.append(link)

            #writing to file .docx:
            write_to_file(document, cur_index, name, link) 

            #increasing cur_index:
            cur_index += 1
    
        
    num_results = len(job_names)
    
    return num_results, job_names, job_links


def beautify_text(text):

    text = re.sub(r'(\n)', '', text)
    text = re.sub(r'[^a-zA-Z\s-]', '', text)
    text = re.sub(r'^(\s)+', '', text)
    
    return text


def write_to_file(document, cur_index, name, link): 

    p = document.add_paragraph()
    add_hyperlink(p, link, str(cur_index) + ' .' + name)  
    #for i in range(num_results):
        #f.write('{}. {: <30} {: <} \n'.format(i+cur_index, job_names[i], job_links[i]))
      

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink



    

document = Document()
cur_index = 1

indeed_job_search(3, 'ds', document, cur_index)
print(cur_index)
drushim_job_search(3, 'ds', document, cur_index)


document.save(r'D:\different_work_projects\test.docx')





    
  